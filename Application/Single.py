import random
from tkinter import *
from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT





def fill_qno_columns(input_docx):
    # Load the output document
    doc = Document(input_docx)
    
    # Access the 0th indexed table (Part-A)
    part_a_table = doc.tables[0]
    
    # Fill the "Q.No." column in Part-A with numbers from 1 to 10
    for index, row in enumerate(part_a_table.rows[1:], start=1):
        row.cells[0].text = str(index)
    
    # Target the Part-B table at index 1
    part_b_table = doc.tables[1]
    
    # Initialize question number for Part-B
    question_number = 11
    # Alternate between numbers and alphabets (a, b)
    alphabet = 'a'
    # Fill the "Q.No." column in Part-B with numbers and alphabets for consecutive questions
    for row in part_b_table.rows[1:]:
        if alphabet == 'a':
            row.cells[0].text = f"{question_number}. {alphabet}."
            alphabet = 'b'
        else:
            row.cells[0].text = f"{question_number}. {alphabet}."
            # Add "(or)" between questions with the same question number
            if alphabet == 'b':
                row.cells[1].text += ""
            question_number += 1
            alphabet = 'a'
    
    # Save the modified document
    output_docx = input_docx.split('.')[0] + '_filled.docx'
    doc.save(output_docx)
    
    print(f"Filled question numbers saved to '{output_docx}'")


def fill_qno_columns_2021(input_docx):
    # Load the output document
    doc = Document(input_docx)
    
    # Access the 0th indexed table (Part-A)
    part_a_table = doc.tables[0]
    
    # Fill the "Q.No." column in Part-A with numbers from 1 to 10
    for index, row in enumerate(part_a_table.rows[1:], start=1):
        row.cells[0].text = str(index)
    
    # Target the Part-B table at index 1
    part_b_table = doc.tables[1]
    
    # Initialize question number for Part-B
    question_number = 11
    # Alternate between numbers and alphabets (a, b)
    alphabet = 'a'
    # Fill the "Q.No." column in Part-B with numbers and alphabets for consecutive questions
    for row in part_b_table.rows[1:]:
        if alphabet == 'a':
            row.cells[0].text = f"{question_number}. {alphabet}."
            alphabet = 'b'
        else:
            row.cells[0].text = f"{question_number}. {alphabet}."
            # Add "(or)" between questions with the same question number
            if alphabet == 'b':
                row.cells[1].text += ""
            question_number += 1
            alphabet = 'a'

    # Target the Part-B table at index 1
    part_b_table = doc.tables[2]

    # Initialize question number for Part-C
    question_number = 16
    # Alternate between numbers and alphabets (a, b)
    alphabet = 'a'
    # Fill the "Q.No." column in Part-B with numbers and alphabets for consecutive questions
    for row in part_b_table.rows[1:]:
        if alphabet == 'a':
            row.cells[0].text = f"{question_number}. {alphabet}."
            alphabet = 'b'
        else:
            row.cells[0].text = f"{question_number}. {alphabet}."
            # Add "(or)" between questions with the same question number
            if alphabet == 'b':
                row.cells[1].text += ""
            question_number += 1
            alphabet = 'a'
    
    # Save the modified document
    output_docx = input_docx.split('.')[0] + '_filled.docx'
    doc.save(output_docx)
    
    print(f"Filled question numbers saved to '{output_docx}'")


def fill_qno_ca1(input_docx):
    # Load the output document
    doc = Document(input_docx)
    
    # Access the 0th indexed table (Part-A)
    part_a_table = doc.tables[0]
    
    # Fill the "Q.No." column in Part-A with numbers from 1 to 10
    for index, row in enumerate(part_a_table.rows[1:], start=1):
        row.cells[0].text = str(index)
    
    # Target the Part-B table at index 1
    part_b_table = doc.tables[1]
    
    # Initialize question number for Part-B
    question_number = 6
    # Alternate between numbers and alphabets (a, b)
    alphabet = 'a'
    # Fill the "Q.No." column in Part-B with numbers and alphabets for consecutive questions
    for row in part_b_table.rows[1:]:
        if alphabet == 'a':
            row.cells[0].text = str(question_number) + '. ' + alphabet + '.'
            alphabet = 'b'
        else:
            row.cells[0].text = str(question_number) + '. ' + alphabet + '.'
            question_number += 1
            alphabet = 'a'
    
    # Save the modified document
    output_docx = input_docx.split('.')[0] + '_filled.docx'
    doc.save(output_docx)
    
    print(f"Filled question numbers saved to '{output_docx}'")


def fill_qno_ca1_2021(input_docx):
    # Load the output document
    doc = Document(input_docx)
    
    # Access the 0th indexed table (Part-A)
    part_a_table = doc.tables[0]
    
    # Fill the "Q.No." column in Part-A with numbers from 1 to 10
    for index, row in enumerate(part_a_table.rows[1:], start=1):
        row.cells[0].text = str(index)
    
    # Target the Part-B table at index 1
    part_b_table = doc.tables[1]
    
    # Initialize question number for Part-B
    question_number = 6
    # Alternate between numbers and alphabets (a, b)
    alphabet = 'a'
    # Fill the "Q.No." column in Part-B with numbers and alphabets for consecutive questions
    for row in part_b_table.rows[1:]:
        if alphabet == 'a':
            row.cells[0].text = str(question_number) + '. ' + alphabet + '.'
            alphabet = 'b'
        else:
            row.cells[0].text = str(question_number) + '. ' + alphabet + '.'
            question_number += 1
            alphabet = 'a'

    # Target the Part-C table at index 1
    part_c_table = doc.tables[2]
    
    # Initialize question number for Part-B
    question_number = 8
    # Alternate between numbers and alphabets (a, b)
    alphabet = 'a'
    # Fill the "Q.No." column in Part-B with numbers and alphabets for consecutive questions
    for row in part_c_table.rows[1:]:
        if alphabet == 'a':
            row.cells[0].text = str(question_number) + '. ' + alphabet + '.'
            alphabet = 'b'
        else:
            row.cells[0].text = str(question_number) + '. ' + alphabet + '.'
            question_number += 1
            alphabet = 'a'
    
    # Save the modified document
    output_docx = input_docx.split('.')[0] + '_filled.docx'
    doc.save(output_docx)
    
    print(f"Filled question numbers saved to '{output_docx}'")


def sem_paper(input_docx, output_docx,Paper_code,subject_name,subject_code,Month_Year, Semester):
    # Load the input Word document
    selected_questions_part_a = []
    selected_questions_part_b = []

    doc = Document(input_docx)
    
    # Extracting data from the tables
    for i in range(0,10):  # Iterate through all tables
        table = doc.tables[i]
        data = []
        for row in table.rows[1:]:  # Skip the header row
            row_data = [cell.text.strip() for cell in row.cells]
            data.append(row_data)
        
        # Shuffle the questions
        random.shuffle(data)

        # Select questions based on the table
        if i % 2 == 0:  # Tables at even indices belong to Part-A
            selected_questions_part_a.extend(data[:2])
        else:           # Tables at odd indices belong to Part-B
            selected_questions_part_b.extend(data[:2])

    # Create a new document to store selected questions
    new_doc = Document()
    
     # Set left and right margins to a smaller value (e.g., 1 cm)
    sections = new_doc.sections
    for section in sections:
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    # Add the first image aligned to the left
    paragraph = new_doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = paragraph.add_run()
    run.add_picture('./assets/Picture 1.jpg', width=Cm(5.66), height=Cm(1.88))

    # Add a tab character to separate the first and second images
    run.add_text('\t\t\t')

    # Add the second image aligned to the right
    run.add_picture('./assets/reg.png', width=Cm(10), height=Cm(1.80))

    # Add document content with font size set to 20 points
    new_doc.add_paragraph(f"Question Paper Code: {Paper_code}", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    new_doc.add_paragraph(f"B.E./B.TECH. DEGREE EXAMINATIONS, {Month_Year}", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    new_doc.add_paragraph("Continuous Assessment II", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    new_doc.add_paragraph(f"{Semester} Semester", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    new_doc.add_paragraph(f"{subject_code} - {subject_name}", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Set font size to 20 points for all paragraphs
    for paragraph in new_doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(18)

    # Add Part-A header
    part_a_heading = new_doc.add_paragraph("Part-A (5 X 2 = 10 Marks)", style='Heading1')
    part_a_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    part_a_heading.runs[0].bold = True
    part_a_heading.runs[0].font.size = Pt(15)

    # Add table for Part-A questions
    new_table_a = new_doc.add_table(rows=1, cols=4)
    new_table_a.style = 'Table Grid'  # Apply table style
    hdr_cells_a = new_table_a.rows[0].cells
    hdr_cells_a[0].text = "Q.No"  # Add Question Number header
    hdr_cells_a[1].text = "Questions"  # Add Questions header
    hdr_cells_a[2].text = "CO’s"  # Add CO's header
    hdr_cells_a[3].text = "Bloom’s Level"  # Add Bloom's Level header

    # Set column width
    new_table_a.columns[1].width = Pt(500)
    new_table_a.columns[0].width = Pt(20)
    new_table_a.columns[2].width = Pt(20)
    new_table_a.columns[3].width = Pt(20)

    for row in new_table_a.rows:
     for cell in row.cells:
         for paragraph in cell.paragraphs:
             for run in paragraph.runs:
                 run.font.size = Pt(14)


    for row_data in selected_questions_part_a:
        row_cells = new_table_a.add_row().cells
        for i, cell_value in enumerate(row_data):
            row_cells[i].text = cell_value

    # Add a paragraph break between Part-A and Part-B
    paragraph = new_doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run()
    run.add_break()

    # Add Part-B header
    part_b_heading = new_doc.add_paragraph("Part-B (2 x 15 = 30 Marks)", style='Heading1')
    part_b_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    part_b_heading.runs[0].bold = True
    part_b_heading.runs[0].font.size = Pt(14)

    # Add table for Part-B questions
    new_table_b = new_doc.add_table(rows=1, cols=4)
    new_table_b.style = 'Table Grid'  # Apply table style
    hdr_cells_b = new_table_b.rows[0].cells
    hdr_cells_b[0].text = "Q.No"  # Add Question Number header
    hdr_cells_b[1].text = "Questions"  # Add Questions header
    hdr_cells_b[2].text = "CO’s"  # Add CO's header
    hdr_cells_b[3].text = "Bloom’s Level"  # Add Bloom's Level header

    #setting width:
    new_table_b.columns[1].width = Pt(500)
    new_table_b.columns[0].width = Pt(20)
    new_table_b.columns[2].width = Pt(20)
    new_table_b.columns[3].width = Pt(20)

    for row_data in selected_questions_part_b:
        row_cells = new_table_b.add_row().cells
        for i, cell_value in enumerate(row_data):
            row_cells[i].text = cell_value
    # Save the new document
    new_doc.save(output_docx)
    print(f"Random questions saved to '{output_docx}'")
    fill_qno_columns("output_questions.docx")


def sem_paper_2021(input_docx, output_docx,Paper_code,subject_name,subject_code,Month_Year, Semester):
    # Load the input Word document
    selected_questions_part_a = []
    selected_questions_part_b = []
    selected_questions_part_c = []
    count=0

    doc = Document(input_docx)
    
    # Extracting data from the tables
    for i in range(0,15):  # Iterate through all tables
        table = doc.tables[i]
        data = []
        for row in table.rows[1:]:  # Skip the header row
            row_data = [cell.text.strip() for cell in row.cells]
            data.append(row_data)
        
        # Shuffle the questions
        random.shuffle(data)

        # Select questions based on the table
        if i == 0 or i == 3 or i == 6 or i == 9 or i == 12:  # Tables at even indices belong to Part-A
            selected_questions_part_a.extend(data[:2])
        elif(i == 1 or i == 4 or i == 7 or i == 10 or i == 13):           # Tables at odd indices belong to Part-B
            selected_questions_part_b.extend(data[:2])
        # elif(i == 2 or i == 4 or i == 7 or i == 10 or i == 13):           # Tables at odd indices belong to Part-B
        else:
            random_bool = random.choice([True, False])
            if(random_bool):
                if(count<2):
                    selected_questions_part_c.extend(data[:1])
                    count +=1

    # Create a new document to store selected questions
    new_doc = Document()
    
     # Set left and right margins to a smaller value (e.g., 1 cm)
    sections = new_doc.sections
    for section in sections:
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    # Add the first image aligned to the left
    paragraph = new_doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = paragraph.add_run()
    run.add_picture('./assets/Picture 1.jpg', width=Cm(5.66), height=Cm(1.88))

    # Add a tab character to separate the first and second images
    run.add_text('\t\t\t')

    # Add the second image aligned to the right
    run.add_picture('./assets/reg.png', width=Cm(10), height=Cm(1.80))

    # Add document content with font size set to 20 points
    new_doc.add_paragraph(f"Question Paper Code: {Paper_code}", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    new_doc.add_paragraph(f"B.E./B.TECH. DEGREE EXAMINATIONS, {Month_Year}", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    new_doc.add_paragraph("Continuous Assessment II", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    new_doc.add_paragraph(f"{Semester} Semester", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    new_doc.add_paragraph(f"{subject_code} - {subject_name}", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Set font size to 20 points for all paragraphs
    for paragraph in new_doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(18)

    # Add Part-A header
    part_a_heading = new_doc.add_paragraph("Part-A (10 X 2 = 20 Marks)", style='Heading1')
    part_a_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    part_a_heading.runs[0].bold = True
    part_a_heading.runs[0].font.size = Pt(15)

    # Add table for Part-A questions
    new_table_a = new_doc.add_table(rows=1, cols=4)
    new_table_a.style = 'Table Grid'  # Apply table style
    hdr_cells_a = new_table_a.rows[0].cells
    hdr_cells_a[0].text = "Q.No"  # Add Question Number header
    hdr_cells_a[1].text = "Questions"  # Add Questions header
    hdr_cells_a[2].text = "CO’s"  # Add CO's header
    hdr_cells_a[3].text = "Bloom’s Level"  # Add Bloom's Level header

    # Set column width
    new_table_a.columns[1].width = Pt(500)
    new_table_a.columns[0].width = Pt(20)
    new_table_a.columns[2].width = Pt(20)
    new_table_a.columns[3].width = Pt(20)

    for row in new_table_a.rows:
     for cell in row.cells:
         for paragraph in cell.paragraphs:
             for run in paragraph.runs:
                 run.font.size = Pt(14)


    for row_data in selected_questions_part_a:
        row_cells = new_table_a.add_row().cells
        for i, cell_value in enumerate(row_data):
            row_cells[i].text = cell_value

    # Add a paragraph break between Part-A and Part-B
    paragraph = new_doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run()
    run.add_break()

    # Add Part-B header
    part_b_heading = new_doc.add_paragraph("Part-B (5 x 13 = 65 Marks)", style='Heading1')
    part_b_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    part_b_heading.runs[0].bold = True
    part_b_heading.runs[0].font.size = Pt(14)

    # Add table for Part-B questions
    new_table_b = new_doc.add_table(rows=1, cols=4)
    new_table_b.style = 'Table Grid'  # Apply table style
    hdr_cells_b = new_table_b.rows[0].cells
    hdr_cells_b[0].text = "Q.No"  # Add Question Number header
    hdr_cells_b[1].text = "Questions"  # Add Questions header
    hdr_cells_b[2].text = "CO’s"  # Add CO's header
    hdr_cells_b[3].text = "Bloom’s Level"  # Add Bloom's Level header

    #setting width:
    new_table_b.columns[1].width = Pt(500)
    new_table_b.columns[0].width = Pt(20)
    new_table_b.columns[2].width = Pt(20)
    new_table_b.columns[3].width = Pt(20)

    for row_data in selected_questions_part_b:
        row_cells = new_table_b.add_row().cells
        for i, cell_value in enumerate(row_data):
            row_cells[i].text = cell_value

    # Add Part-C header
    part_c_heading = new_doc.add_paragraph("Part-C (1 x 15 = 15 Marks)", style='Heading1')
    part_c_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    part_c_heading.runs[0].bold = True
    part_c_heading.runs[0].font.size = Pt(14)

    # Add table for Part-B questions
    new_table_c = new_doc.add_table(rows=1, cols=4)
    new_table_c.style = 'Table Grid'  # Apply table style
    hdr_cells_c = new_table_b.rows[0].cells
    hdr_cells_c[0].text = "Q.No"  # Add Question Number header
    hdr_cells_c[1].text = "Questions"  # Add Questions header
    hdr_cells_c[2].text = "CO’s"  # Add CO's header
    hdr_cells_c[3].text = "Bloom’s Level"  # Add Bloom's Level header

    #setting width:
    new_table_c.columns[1].width = Pt(500)
    new_table_c.columns[0].width = Pt(20)
    new_table_c.columns[2].width = Pt(20)
    new_table_c.columns[3].width = Pt(20)

    for row_data in selected_questions_part_c:
        row_cells = new_table_c.add_row().cells
        for i, cell_value in enumerate(row_data):
            row_cells[i].text = cell_value

    # Save the new document
    new_doc.save(output_docx)
    print(f"Random questions saved to '{output_docx}'")
    fill_qno_columns_2021("output_questions.docx")




from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def ca2_paper(input_docx, output_docx, Paper_code, subject_name, subject_code, num, Month_Year, Semester,start,end):
    # Load the input Word document
    selected_questions_part_a = []
    selected_questions_part_b = []
    count = 0
    doc = Document(input_docx)
    rem = 5 - int(num)

    # Extracting data from the tables
    for i in range(start, start+2):  # Iterate through all tables
        table = doc.tables[i]
        data = []
        for row in table.rows[1:]:  # Skip the header row
            row_data = [cell.text.strip() for cell in row.cells]
            data.append(row_data)

        # Shuffle the questions
        random.shuffle(data)

        # Select questions based on the table
        if i % 2 == 0:  # Tables at even indices belong to Part-A
            if count == 0:
                selected_questions_part_a.extend(data[:num])
                count = 1
            else:
                selected_questions_part_a.extend(data[:rem])
        else:  # Tables at odd indices belong to Part-B
            selected_questions_part_b.extend(data[:2])

    for i in range(end, end+2):  # Iterate through all tables
        table = doc.tables[i]
        data = []
        for row in table.rows[1:]:  # Skip the header row
            row_data = [cell.text.strip() for cell in row.cells]
            data.append(row_data)

        # Shuffle the questions
        random.shuffle(data)

        # Select questions based on the table
        if i % 2 == 0:  # Tables at even indices belong to Part-A
            if count == 0:
                selected_questions_part_a.extend(data[:num])
                count = 1
            else:
                selected_questions_part_a.extend(data[:rem])
        else:  # Tables at odd indices belong to Part-B
            selected_questions_part_b.extend(data[:2])

    # Create a new document to store selected questions
    new_doc = Document()

    # Set left and right margins to a smaller value (e.g., 1 cm)
    sections = new_doc.sections
    for section in sections:
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    # Add the first image aligned to the left
    paragraph = new_doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = paragraph.add_run()
    run.add_picture('./assets/Picture 1.jpg', width=Cm(5.66), height=Cm(1.88))

    # Add a tab character to separate the first and second images
    run.add_text('\t\t\t')

    # Add the second image aligned to the right
    run.add_picture('./assets/reg.png', width=Cm(10), height=Cm(1.80))

    # Add document content with font size set to 20 points
    new_doc.add_paragraph(f"Question Paper Code: {Paper_code}", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    new_doc.add_paragraph(f"B.E./B.TECH. DEGREE EXAMINATIONS, {Month_Year}", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    new_doc.add_paragraph("Continuous Assessment II", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    new_doc.add_paragraph(f"{Semester} Semester", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    new_doc.add_paragraph(f"{subject_code} - {subject_name}", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Set font size to 20 points for all paragraphs
    for paragraph in new_doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(18)

    # Add Part-A header
    part_a_heading = new_doc.add_paragraph("Part-A (5 X 2 = 10 Marks)", style='Heading1')
    part_a_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    part_a_heading.runs[0].bold = True
    part_a_heading.runs[0].font.size = Pt(15)

    # Add table for Part-A questions
    new_table_a = new_doc.add_table(rows=1, cols=4)
    new_table_a.style = 'Table Grid'  # Apply table style
    hdr_cells_a = new_table_a.rows[0].cells
    hdr_cells_a[0].text = "Q.No"  # Add Question Number header
    hdr_cells_a[1].text = "Questions"  # Add Questions header
    hdr_cells_a[2].text = "CO’s"  # Add CO's header
    hdr_cells_a[3].text = "Bloom’s Level"  # Add Bloom's Level header

    # Set column width
    new_table_a.columns[1].width = Pt(500)
    new_table_a.columns[0].width = Pt(50)
    new_table_a.columns[2].width = Pt(50)
    new_table_a.columns[3].width = Pt(50)

    # Set row height and font properties
    for row in new_table_a.rows:
        row.height = Pt(20)
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(11)

    for row_data in selected_questions_part_a:
        row_cells = new_table_a.add_row().cells
        for i, cell_value in enumerate(row_data):
            row_cells[i].text = cell_value

    # Add a paragraph break between Part-A and Part-B
    paragraph = new_doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run()
    run.add_break()

    # Add Part-B header
    part_b_heading = new_doc.add_paragraph("Part-B (2 x 15 = 30 Marks)", style='Heading1')
    part_b_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    part_b_heading.runs[0].bold = True
    part_b_heading.runs[0].font.size = Pt(14)

    # Add table for Part-B questions
    new_table_b = new_doc.add_table(rows=1, cols=4)
    new_table_b.style = 'Table Grid'  # Apply table style
    hdr_cells_b = new_table_b.rows[0].cells
    hdr_cells_b[0].text = "Q.No"  # Add Question Number header
    hdr_cells_b[1].text = "Questions"  # Add Questions header
    hdr_cells_b[2].text = "CO’s"  # Add CO's header
    hdr_cells_b[3].text = "Bloom’s Level"  # Add Bloom's Level header

    # Set column width
    new_table_b.columns[1].width = Pt(500)
    new_table_b.columns[0].width = Pt(50)
    new_table_b.columns[2].width = Pt(50)
    new_table_b.columns[3].width = Pt(50)

    # Set row height and font properties
    for row in new_table_b.rows:
        row.height = Pt(20)
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(11)

    for row_data in selected_questions_part_b:
        row_cells = new_table_b.add_row().cells
        for i, cell_value in enumerate(row_data):
            row_cells[i].text = cell_value

    # Save the new document
    new_doc.save(output_docx)
    print(f"Random questions saved to '{output_docx}'")
    fill_qno_ca1("output_questions.docx")





#2021 CA


def ca2_paper_2021(input_docx, output_docx, Paper_code, subject_name, subject_code, num, Month_Year, Semester,start,end):
    # Load the input Word document
    selected_questions_part_a = []
    selected_questions_part_b = []
    selected_questions_part_c = []
    count = 0
    doc = Document(input_docx)
    rem = 5 - int(num)

    # Extracting data from the tables
    for i in range(start, start+3):  # Iterate through all tables
        table = doc.tables[i]
        data = []
        for row in table.rows[1:]:  # Skip the header row
            row_data = [cell.text.strip() for cell in row.cells]
            data.append(row_data)

        # Shuffle the questions
        random.shuffle(data)

        # Select questions based on the table
        if i == 0 or i == 3 or i == 6 or i == 9 or i == 12 : # Tables at even indices belong to Part-A
            if count == 0:
                selected_questions_part_a.extend(data[:num])
                count = 1
            else:
                selected_questions_part_a.extend(data[:rem])
        elif(i == 1 or i == 4 or i == 7 or i == 10 or i == 13):  # Tables at odd indices belong to Part-B
            selected_questions_part_b.extend(data[:2])

        else:
            selected_questions_part_c.extend(data[:1])

    for i in range(end, end+3):  # Iterate through all tables
        table = doc.tables[i]
        data = []
        for row in table.rows[1:]:  # Skip the header row
            row_data = [cell.text.strip() for cell in row.cells]
            data.append(row_data)

        # Shuffle the questions
        random.shuffle(data)

         # Select questions based on the table
        if i == 0 or i == 3 or i == 6 or i == 9 or i == 12 : # Tables at even indices belong to Part-A
            if count == 0:
                selected_questions_part_a.extend(data[:num])
                count = 1
            else:
                selected_questions_part_a.extend(data[:rem])
        elif(i == 1 or i == 4 or i == 7 or i == 10 or i == 13):  # Tables at odd indices belong to Part-B
            selected_questions_part_b.extend(data[:2])

        else:
            selected_questions_part_c.extend(data[:1])

    # Create a new document to store selected questions
    new_doc = Document()

    # Set left and right margins to a smaller value (e.g., 1 cm)
    sections = new_doc.sections
    for section in sections:
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    # Add the first image aligned to the left
    paragraph = new_doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = paragraph.add_run()
    run.add_picture('./assets/Picture 1.jpg', width=Cm(5.66), height=Cm(1.88))

    # Add a tab character to separate the first and second images
    run.add_text('\t\t\t')

    # Add the second image aligned to the right
    run.add_picture('./assets/reg.png', width=Cm(10), height=Cm(1.80))

    # Add document content with font size set to 20 points
    new_doc.add_paragraph(f"Question Paper Code: {Paper_code}", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    new_doc.add_paragraph(f"B.E./B.TECH. DEGREE EXAMINATIONS, {Month_Year}", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    new_doc.add_paragraph("Continuous Assessment II", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    new_doc.add_paragraph(f"{Semester} Semester", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    new_doc.add_paragraph(f"{subject_code} - {subject_name}", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Set font size to 20 points for all paragraphs
    for paragraph in new_doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(18)

    # Add Part-A header
    part_a_heading = new_doc.add_paragraph("Part-A (5 X 2 = 10 Marks)", style='Heading1')
    part_a_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    part_a_heading.runs[0].bold = True
    part_a_heading.runs[0].font.size = Pt(15)

    # Add table for Part-A questions
    new_table_a = new_doc.add_table(rows=1, cols=4)
    new_table_a.style = 'Table Grid'  # Apply table style
    hdr_cells_a = new_table_a.rows[0].cells
    hdr_cells_a[0].text = "Q.No"  # Add Question Number header
    hdr_cells_a[1].text = "Questions"  # Add Questions header
    hdr_cells_a[2].text = "CO’s"  # Add CO's header
    hdr_cells_a[3].text = "Bloom’s Level"  # Add Bloom's Level header

    # Set column width
    new_table_a.columns[1].width = Pt(500)
    new_table_a.columns[0].width = Pt(50)
    new_table_a.columns[2].width = Pt(50)
    new_table_a.columns[3].width = Pt(50)

    # Set row height and font properties
    for row in new_table_a.rows:
        row.height = Pt(20)
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(11)

    for row_data in selected_questions_part_a:
        row_cells = new_table_a.add_row().cells
        for i, cell_value in enumerate(row_data):
            row_cells[i].text = cell_value

    # Add a paragraph break between Part-A and Part-B
    paragraph = new_doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run()
    run.add_break()

    # Add Part-B header
    part_b_heading = new_doc.add_paragraph("Part-B (2 x 13 = 26 Marks)", style='Heading1')
    part_b_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    part_b_heading.runs[0].bold = True
    part_b_heading.runs[0].font.size = Pt(14)

    # Add table for Part-B questions
    new_table_b = new_doc.add_table(rows=1, cols=4)
    new_table_b.style = 'Table Grid'  # Apply table style
    hdr_cells_b = new_table_b.rows[0].cells
    hdr_cells_b[0].text = "Q.No"  # Add Question Number header
    hdr_cells_b[1].text = "Questions"  # Add Questions header
    hdr_cells_b[2].text = "CO’s"  # Add CO's header
    hdr_cells_b[3].text = "Bloom’s Level"  # Add Bloom's Level header

    # Set column width
    new_table_b.columns[1].width = Pt(500)
    new_table_b.columns[0].width = Pt(50)
    new_table_b.columns[2].width = Pt(50)
    new_table_b.columns[3].width = Pt(50)

    # Set row height and font properties
    for row in new_table_b.rows:
        row.height = Pt(20)
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(11)

    for row_data in selected_questions_part_b:
        row_cells = new_table_b.add_row().cells
        for i, cell_value in enumerate(row_data):
            row_cells[i].text = cell_value

     # Add a paragraph break between Part-A and Part-B
    paragraph = new_doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run()
    run.add_break()


    # Add Part-C header
    part_c_heading = new_doc.add_paragraph("Part-C (1 x 14 = 14 Marks)", style='Heading1')
    part_c_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    part_c_heading.runs[0].bold = True
    part_c_heading.runs[0].font.size = Pt(14)

    # Add table for Part-B questions
    new_table_c = new_doc.add_table(rows=1, cols=4)
    new_table_c.style = 'Table Grid'  # Apply table style
    hdr_cells_c = new_table_b.rows[0].cells
    hdr_cells_c[0].text = "Q.No"  # Add Question Number header
    hdr_cells_c[1].text = "Questions"  # Add Questions header
    hdr_cells_c[2].text = "CO’s"  # Add CO's header
    hdr_cells_c[3].text = "Bloom’s Level"  # Add Bloom's Level header

    # Set column width
    new_table_c.columns[1].width = Pt(500)
    new_table_c.columns[0].width = Pt(50)
    new_table_c.columns[2].width = Pt(50)
    new_table_c.columns[3].width = Pt(50)

    # Set row height and font properties
     # Set row height and font properties
    for row in new_table_b.rows:
        row.height = Pt(20)
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(11)

    for row_data in selected_questions_part_c:
        row_cells = new_table_c.add_row().cells
        for i, cell_value in enumerate(row_data):
            row_cells[i].text = cell_value


    # Save the new document
    new_doc.save(output_docx)
    print(f"Random questions saved to '{output_docx}'")
    fill_qno_ca1_2021("output_questions.docx")


# ca2_paper("Format.docx", "output_questions.docx", "P123", "Physics", "PHY101", 3, "March 2024", "3rd")
# ca2_paper_2021("Format.docx", "output_questions.docx", "P123", "Physics", "PHY101", 3, "March 2024", "3rd")
# sem_paper("Format.docx", "output_questions.docx", "P123", "Physics", "PHY101","March 2024", "3rd")
# sem_paper_2021("Format_2021.docx", "output_questions.docx", "P123", "Physics", "PHY101","March 2024", "3rd")






#UI---------------------------------------------








#=====================================UI---------------------------------------------


# from main import sem_paper, ca2_paper
from tkinter import *
from tkinter import filedialog
import os

# Create a Tkinter instance
root = Tk()

# Set the window title
root.title("Question Paper Generator")

# Create a Canvas widget
canvas = Canvas(root)
canvas.pack(side=LEFT, fill="both", expand=True)

# Add a scrollbar
scrollbar = Scrollbar(root, orient="vertical", command=canvas.yview)
scrollbar.pack(side=RIGHT, fill="y")

# Configure the canvas
canvas.configure(yscrollcommand=scrollbar.set)
canvas.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))

# Create a frame inside the canvas to hold the widgets
frame = Frame(canvas)

# Add the frame to the canvas
canvas.create_window((0, 0), window=frame, anchor="nw")

# Global variables to store inputs
subject_name = ""
sub_code = ""
Paper_code = ""
no_of_questions = 0
output_docx_path = ""
output_docx_filled_path = ""
global start, end
start = -1
end = -1
strt=1
ed=-1
# Function to update global variables with input values
def Unit1():
    global start, end,strt,ed
    if start == -1:
        start = 0
    else:
        end = 0

    if strt == -1:
        strt = 0
    else:
        ed = 0


# Function to update global variables with input values
def Unit2():
    global start, end,strt,ed
    if start == -1:
        start = 2
    else:
        end = 2

    if strt == -1:
        strt = 3
    else:
        ed = 3

# Function to update global variables with input values
def Unit3():
    global start, end,strt,ed
    if start == -1:
        start = 4
    else:
        end = 4

    if strt == -1:
        strt = 6
    else:
        ed = 6

# Function to update global variables with input values
def Unit4():
    global start, end,strt,ed
    if start == -1:
        start = 6
    else:
        end = 6
    
    if strt == -1:
        strt = 9
    else:
        ed = 9

# Function to update global variables with input values
def Unit5():
    global start, end,strt,ed
    if start == -1:
        start = 8
    else:
        end = 8

    if strt == -1:
        strt = 12
    else:
        ed = 12

# Function to get user inputs
def get_inputs():
    global subject_name, sub_code, Paper_code, weightage, MonthYear, semester
    subject_name = subject_name_entry.get()
    sub_code = sub_code_entry.get()
    Paper_code = Paper_code_entry.get()
    weightage = int(no_of_questions.get())
    MonthYear = Month_Year.get()
    semester = Semester.get()

# Callback function to open file dialog
def openFile():
    global path
    path = filedialog.askopenfilename()
    print(path)

# Callback function to generate question paper for Semester/Model
def sem_button_callback():
    get_inputs()
    sem_paper(path, "output_questions.docx", Paper_code, subject_name, sub_code, MonthYear, semester)
    sem_paper(path, "output_questions_filled.docx", Paper_code, subject_name, sub_code, MonthYear, semester)
    global output_docx_path, output_docx_filled_path
    output_docx_path = "output_questions.docx"
    output_docx_filled_path = "output_questions_filled.docx"
    open_generated_docx()

def sem_2021_callback():
    get_inputs()
    sem_paper_2021(path, "output_questions.docx", Paper_code, subject_name, sub_code, MonthYear, semester)
    sem_paper_2021(path, "output_questions_filled.docx", Paper_code, subject_name, sub_code, MonthYear, semester)
    global output_docx_path, output_docx_filled_path
    output_docx_path = "output_questions.docx"
    output_docx_filled_path = "output_questions_filled.docx"
    open_generated_docx()


# Callback function to generate question paper for CA2
def ca2_paper_callback():
    get_inputs()
    ca2_paper(path, "output_questions.docx", Paper_code, subject_name, sub_code, weightage, Month_Year, semester, start, end)
    ca2_paper(path, "output_questions_filled.docx", Paper_code, subject_name, sub_code, weightage, Month_Year, semester, start, end)
    global output_docx_path, output_docx_filled_path
    output_docx_path = "output_questions.docx"
    output_docx_filled_path = "output_questions_filled.docx"
    open_generated_docx()


def ca2_paper_2021_callback():
    get_inputs()
    ca2_paper_2021(path, "output_questions.docx", Paper_code, subject_name, sub_code, weightage, Month_Year, semester, strt, ed)
    ca2_paper_2021(path, "output_questions_filled.docx", Paper_code, subject_name, sub_code, weightage, Month_Year, semester, strt, ed)
    global output_docx_path, output_docx_filled_path
    output_docx_path = "output_questions.docx"
    output_docx_filled_path = "output_questions_filled.docx"
    open_generated_docx()

# Function to open the generated documents
def open_generated_docx():
    if output_docx_path:
        os.system(f'open "{output_docx_path}"')
    if output_docx_filled_path:
        os.system(f'open "{output_docx_filled_path}"')

# Create and pack the widgets
label = Label(frame, text='Select the formatted question bank', font=("Helvetica", 18))
label.pack(pady=10)

button = Button(frame, text='Open File', command=openFile)
button.pack(pady=10)

subject_name_label = Label(frame, text='Enter Subject Name:', font=("Helvetica", 14))
subject_name_label.pack(pady=5)
subject_name_entry = Entry(frame, font=("Helvetica", 14))
subject_name_entry.pack(pady=5)

sub_code_label = Label(frame, text='Enter Subject Code:', font=("Helvetica", 14))
sub_code_label.pack(pady=5)
sub_code_entry = Entry(frame, font=("Helvetica", 14))
sub_code_entry.pack(pady=5)

Paper_code_label = Label(frame, text='Enter Paper Code:', font=("Helvetica", 14))
Paper_code_label.pack(pady=5)
Paper_code_entry = Entry(frame, font=("Helvetica", 14))
Paper_code_entry.pack(pady=5)

no_of_questions_label = Label(frame, text='Enter weightage of Part-A:', font=("Helvetica", 14))
no_of_questions_label.pack(pady=5)
no_of_questions = Entry(frame, font=("Helvetica", 14))
no_of_questions.pack(pady=5)

Month_Year_label = Label(frame, text='Month & Year:', font=("Helvetica", 14))
Month_Year_label.pack(pady=5)
Month_Year = Entry(frame, font=("Helvetica", 14))
Month_Year.pack(pady=5)

Semester_label = Label(frame, text='Semester:', font=("Helvetica", 14))
Semester_label.pack(pady=5)
Semester = Entry(frame, font=("Helvetica", 14))
Semester.pack(pady=5)

Unit_buttons_frame = Frame(frame)
Unit_buttons_frame.pack(pady=20)

label = Label(frame, text='Choose unit :', font=("Helvetica", 18))
label.pack(pady=10)

Unit1_button = Button(Unit_buttons_frame, text='Unit 1', font=("Helvetica", 14), command=Unit1)
Unit1_button.pack(side=LEFT, padx=5)

Unit2_button = Button(Unit_buttons_frame, text='Unit 2', font=("Helvetica", 14), command=Unit2)
Unit2_button.pack(side=LEFT, padx=5)

Unit3_button = Button(Unit_buttons_frame, text='Unit 3', font=("Helvetica", 14), command=Unit3)
Unit3_button.pack(side=LEFT, padx=5)

Unit4_button = Button(Unit_buttons_frame, text='Unit 4', font=("Helvetica", 14), command=Unit4)
Unit4_button.pack(side=LEFT, padx=5)

Unit5_button = Button(Unit_buttons_frame, text='Unit 5', font=("Helvetica", 14), command=Unit5)
Unit5_button.pack(side=LEFT, padx=5)

generate_label = Label(frame, text='Which format of the question paper do you want?', font=("Helvetica", 18))
generate_label.pack(pady=10)

sem_button = Button(frame, text='Semester/Model', font=("Helvetica", 16), command=sem_button_callback)
sem_button.pack(pady=10)


sem_button_2021 = Button(frame, text='Semester/Model_2021', font=("Helvetica", 16), command=sem_2021_callback)
sem_button_2021.pack(pady=10)


ca2_button = Button(frame, text='Generate CA question', font=("Helvetica", 16), command=ca2_paper_callback)
ca2_button.pack(pady=10)

ca2_button_2021 = Button(frame, text='Generate CA question 2021', font=("Helvetica", 16), command=ca2_paper_2021_callback)
ca2_button_2021.pack(pady=10)

# Update the scroll region to include the new widgets
frame.update_idletasks()
canvas.config(scrollregion=canvas.bbox("all"))

# Run the Tkinter event loop
root.mainloop()

