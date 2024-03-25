import random
from tkinter import *
from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT




def fill_qno_columns(input_docx):
    # Load the output document
    doc = Document(input_docx)
    
    # Access the 0th indexed table (Part-A)
    part_a_table = doc.tables[0]
    
    # Fill the "Q.No." column in Part-A with numbers from 1 to 10
    for index, row in enumerate(part_a_table.rows[1:], start=1):
        row.cells[0].text = str(index)
    
    # Target the Part-B table at index 1
    part_b_table = doc.tables[1]
    
    # Initialize question number for Part-B
    question_number = 11
    # Alternate between numbers and alphabets (a, b)
    alphabet = 'a'
    # Fill the "Q.No." column in Part-B with numbers and alphabets for consecutive questions
    for row in part_b_table.rows[1:]:
        if alphabet == 'a':
            row.cells[0].text = f"{question_number}. {alphabet}."
            alphabet = 'b'
        else:
            row.cells[0].text = f"{question_number}. {alphabet}."
            # Add "(or)" between questions with the same question number
            if alphabet == 'b':
                row.cells[1].text += ""
            question_number += 1
            alphabet = 'a'
    
    # Save the modified document
    output_docx = input_docx.split('.')[0] + '_filled.docx'
    doc.save(output_docx)
    
    print(f"Filled question numbers saved to '{output_docx}'")

def fill_qno_ca1(input_docx):
    # Load the output document
    doc = Document(input_docx)
    
    # Access the 0th indexed table (Part-A)
    part_a_table = doc.tables[0]
    
    # Fill the "Q.No." column in Part-A with numbers from 1 to 10
    for index, row in enumerate(part_a_table.rows[1:], start=1):
        row.cells[0].text = str(index)
    
    # Target the Part-B table at index 1
    part_b_table = doc.tables[1]
    
    # Initialize question number for Part-B
    question_number = 6
    # Alternate between numbers and alphabets (a, b)
    alphabet = 'a'
    # Fill the "Q.No." column in Part-B with numbers and alphabets for consecutive questions
    for row in part_b_table.rows[1:]:
        if alphabet == 'a':
            row.cells[0].text = str(question_number) + '. ' + alphabet + '.'
            alphabet = 'b'
        else:
            row.cells[0].text = str(question_number) + '. ' + alphabet + '.'
            question_number += 1
            alphabet = 'a'
    
    # Save the modified document
    output_docx = input_docx.split('.')[0] + '_filled.docx'
    doc.save(output_docx)
    
    print(f"Filled question numbers saved to '{output_docx}'")

def sem_paper(input_docx, output_docx,Paper_code,subject_name,subject_code,Month_Year, Semester):
    # Load the input Word document
    selected_questions_part_a = []
    selected_questions_part_b = []

    doc = Document(input_docx)
    
    # Extracting data from the tables
    for i in range(0,10):  # Iterate through all tables
        table = doc.tables[i]
        data = []
        for row in table.rows[1:]:  # Skip the header row
            row_data = [cell.text.strip() for cell in row.cells]
            data.append(row_data)
        
        # Shuffle the questions
        random.shuffle(data)

        # Select questions based on the table
        if i % 2 == 0:  # Tables at even indices belong to Part-A
            selected_questions_part_a.extend(data[:2])
        else:           # Tables at odd indices belong to Part-B
            selected_questions_part_b.extend(data[:2])

    # Create a new document to store selected questions
    new_doc = Document()
    
     # Set left and right margins to a smaller value (e.g., 1 cm)
    sections = new_doc.sections
    for section in sections:
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    # Add the first image aligned to the left
    paragraph = new_doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = paragraph.add_run()
    run.add_picture('./assets/Picture 1.jpg', width=Cm(5.66), height=Cm(1.88))

    # Add a tab character to separate the first and second images
    run.add_text('\t\t\t')

    # Add the second image aligned to the right
    run.add_picture('./assets/reg.png', width=Cm(10), height=Cm(1.80))

    # Add document content with font size set to 20 points
    new_doc.add_paragraph(f"Question Paper Code: {Paper_code}", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    new_doc.add_paragraph(f"B.E./B.TECH. DEGREE EXAMINATIONS, {Month_Year}", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    new_doc.add_paragraph("Continuous Assessment II", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    new_doc.add_paragraph(f"{Semester} Semester", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    new_doc.add_paragraph(f"{subject_code} - {subject_name}", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Set font size to 20 points for all paragraphs
    for paragraph in new_doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(18)

    # Add Part-A header
    part_a_heading = new_doc.add_paragraph("Part-A (5 X 2 = 10 Marks)", style='Heading1')
    part_a_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    part_a_heading.runs[0].bold = True
    part_a_heading.runs[0].font.size = Pt(15)

    # Add table for Part-A questions
    new_table_a = new_doc.add_table(rows=1, cols=4)
    new_table_a.style = 'Table Grid'  # Apply table style
    hdr_cells_a = new_table_a.rows[0].cells
    hdr_cells_a[0].text = "Q.No"  # Add Question Number header
    hdr_cells_a[1].text = "Questions"  # Add Questions header
    hdr_cells_a[2].text = "CO’s"  # Add CO's header
    hdr_cells_a[3].text = "Bloom’s Level"  # Add Bloom's Level header

    # Set column width
    new_table_a.columns[1].width = Pt(500)
    new_table_a.columns[0].width = Pt(20)
    new_table_a.columns[2].width = Pt(20)
    new_table_a.columns[3].width = Pt(20)

    for row in new_table_a.rows:
     for cell in row.cells:
         for paragraph in cell.paragraphs:
             for run in paragraph.runs:
                 run.font.size = Pt(14)


    for row_data in selected_questions_part_a:
        row_cells = new_table_a.add_row().cells
        for i, cell_value in enumerate(row_data):
            row_cells[i].text = cell_value

    # Add a paragraph break between Part-A and Part-B
    paragraph = new_doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run()
    run.add_break()

    # Add Part-B header
    part_b_heading = new_doc.add_paragraph("Part-B (2 x 15 = 30 Marks)", style='Heading1')
    part_b_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    part_b_heading.runs[0].bold = True
    part_b_heading.runs[0].font.size = Pt(14)

    # Add table for Part-B questions
    new_table_b = new_doc.add_table(rows=1, cols=4)
    new_table_b.style = 'Table Grid'  # Apply table style
    hdr_cells_b = new_table_b.rows[0].cells
    hdr_cells_b[0].text = "Q.No"  # Add Question Number header
    hdr_cells_b[1].text = "Questions"  # Add Questions header
    hdr_cells_b[2].text = "CO’s"  # Add CO's header
    hdr_cells_b[3].text = "Bloom’s Level"  # Add Bloom's Level header

    #setting width:
    new_table_b.columns[1].width = Pt(500)
    new_table_b.columns[0].width = Pt(20)
    new_table_b.columns[2].width = Pt(20)
    new_table_b.columns[3].width = Pt(20)

    for row_data in selected_questions_part_b:
        row_cells = new_table_b.add_row().cells
        for i, cell_value in enumerate(row_data):
            row_cells[i].text = cell_value
    # Save the new document
    new_doc.save(output_docx)
    print(f"Random questions saved to '{output_docx}'")
    fill_qno_columns("output_questions.docx")

    
# def ca1_paper(input_docx, output_docx,Paper_code,subject_name,subject_code,num):
#     # Load the input Word document
#     selected_questions_part_a = []
#     selected_questions_part_b = []
#     count=0
#     doc = Document(input_docx)
#     # num = random.choice([2, 3])
#     rem=5-num
#     # Extracting data from the tables
#     for i in range(0,4):  # Iterate through all tables
#         table = doc.tables[i]
#         data = []
#         for row in table.rows[1:]:  # Skip the header row
#             row_data = [cell.text.strip() for cell in row.cells]
#             data.append(row_data)
        
#         # Shuffle the questions
#         random.shuffle(data)

#         # Select questions based on the table
#         if i % 2 == 0:  # Tables at even indices belong to Part-A
#             if(count==0):
#                 selected_questions_part_a.extend(data[:num])
#                 count=1
#             else:
#                 selected_questions_part_a.extend(data[:rem])

#         else:           # Tables at odd indices belong to Part-B
#             selected_questions_part_b.extend(data[:2])

#     # Create a new document to store selected questions
#     new_doc = Document()
    
    
#     # Add heading to the document
#     new_doc.add_heading("K.C.G COLLEGE OF TECHNOLOGY", level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#     new_doc.add_paragraph(f"Question Paper Code: {Paper_code}", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#     new_doc.add_paragraph(f"Continous Assesment I", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#     new_doc.add_paragraph(f"{subject_code} - {subject_name}", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
#     # Add Part-A header
#     new_doc.add_paragraph("Part-A", style='Heading1').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

#     # Add table for Part-A questions
#     new_table_a = new_doc.add_table(rows=1, cols=4)
#     new_table_a.style = 'Table Grid'  # Apply table style
#     hdr_cells_a = new_table_a.rows[0].cells
#     hdr_cells_a[0].text = "Q.No"  # Add Question Number header
#     hdr_cells_a[1].text = "Questions"  # Add Questions header
#     hdr_cells_a[2].text = "CO’s"  # Add CO's header
#     hdr_cells_a[3].text = "Bloom’s Level"  # Add Bloom's Level header

#     #setting Width
#     new_table_a.columns[1].width = Pt(300)

#     for row_data in selected_questions_part_a:
#         row_cells = new_table_a.add_row().cells
#         for i, cell_value in enumerate(row_data):
#             row_cells[i].text = cell_value

#     # Add a paragraph break between Part-A and Part-B
#     paragraph = new_doc.add_paragraph()
#     paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#     run = paragraph.add_run()
#     run.add_break()

#     # Add Part-B header
#     new_doc.add_paragraph("Part-B", style='Heading1').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

#     # Add table for Part-B questions
#     new_table_b = new_doc.add_table(rows=1, cols=4)
#     new_table_b.style = 'Table Grid'  # Apply table style
#     hdr_cells_b = new_table_b.rows[0].cells
#     hdr_cells_b[0].text = "Q.No"  # Add Question Number header
#     hdr_cells_b[1].text = "Questions"  # Add Questions header
#     hdr_cells_b[2].text = "CO’s"  # Add CO's header
#     hdr_cells_b[3].text = "Bloom’s Level"  # Add Bloom's Level header

#     #setting width:
#     new_table_b.columns[1].width = Pt(300)

#     for row_data in selected_questions_part_b:
#         row_cells = new_table_b.add_row().cells
#         for i, cell_value in enumerate(row_data):
#             row_cells[i].text = cell_value
#     # Save the new document
#     new_doc.save(output_docx)
#     print(f"Random questions saved to '{output_docx}'")
#     fill_qno_ca1("output_questions.docx")




from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def ca2_paper(input_docx, output_docx, Paper_code, subject_name, subject_code, num, Month_Year, Semester,start,end):
    # Load the input Word document
    selected_questions_part_a = []
    selected_questions_part_b = []
    count = 0
    doc = Document(input_docx)
    rem = 5 - int(num)

    # Extracting data from the tables
    for i in range(start, start+2):  # Iterate through all tables
        table = doc.tables[i]
        data = []
        for row in table.rows[1:]:  # Skip the header row
            row_data = [cell.text.strip() for cell in row.cells]
            data.append(row_data)

        # Shuffle the questions
        random.shuffle(data)

        # Select questions based on the table
        if i % 2 == 0:  # Tables at even indices belong to Part-A
            if count == 0:
                selected_questions_part_a.extend(data[:num])
                count = 1
            else:
                selected_questions_part_a.extend(data[:rem])
        else:  # Tables at odd indices belong to Part-B
            selected_questions_part_b.extend(data[:2])

    for i in range(end, end+2):  # Iterate through all tables
        table = doc.tables[i]
        data = []
        for row in table.rows[1:]:  # Skip the header row
            row_data = [cell.text.strip() for cell in row.cells]
            data.append(row_data)

        # Shuffle the questions
        random.shuffle(data)

        # Select questions based on the table
        if i % 2 == 0:  # Tables at even indices belong to Part-A
            if count == 0:
                selected_questions_part_a.extend(data[:num])
                count = 1
            else:
                selected_questions_part_a.extend(data[:rem])
        else:  # Tables at odd indices belong to Part-B
            selected_questions_part_b.extend(data[:2])

    # Create a new document to store selected questions
    new_doc = Document()

    # Set left and right margins to a smaller value (e.g., 1 cm)
    sections = new_doc.sections
    for section in sections:
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    # Add the first image aligned to the left
    paragraph = new_doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = paragraph.add_run()
    run.add_picture('./assets/Picture 1.jpg', width=Cm(5.66), height=Cm(1.88))

    # Add a tab character to separate the first and second images
    run.add_text('\t\t\t')

    # Add the second image aligned to the right
    run.add_picture('./assets/reg.png', width=Cm(10), height=Cm(1.80))

    # Add document content with font size set to 20 points
    new_doc.add_paragraph(f"Question Paper Code: {Paper_code}", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    new_doc.add_paragraph(f"B.E./B.TECH. DEGREE EXAMINATIONS, {Month_Year}", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    new_doc.add_paragraph("Continuous Assessment II", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    new_doc.add_paragraph(f"{Semester} Semester", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    new_doc.add_paragraph(f"{subject_code} - {subject_name}", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Set font size to 20 points for all paragraphs
    for paragraph in new_doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(18)

    # Add Part-A header
    part_a_heading = new_doc.add_paragraph("Part-A (5 X 2 = 10 Marks)", style='Heading1')
    part_a_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    part_a_heading.runs[0].bold = True
    part_a_heading.runs[0].font.size = Pt(15)

    # Add table for Part-A questions
    new_table_a = new_doc.add_table(rows=1, cols=4)
    new_table_a.style = 'Table Grid'  # Apply table style
    hdr_cells_a = new_table_a.rows[0].cells
    hdr_cells_a[0].text = "Q.No"  # Add Question Number header
    hdr_cells_a[1].text = "Questions"  # Add Questions header
    hdr_cells_a[2].text = "CO’s"  # Add CO's header
    hdr_cells_a[3].text = "Bloom’s Level"  # Add Bloom's Level header

    # Set column width
    new_table_a.columns[1].width = Pt(500)
    new_table_a.columns[0].width = Pt(50)
    new_table_a.columns[2].width = Pt(50)
    new_table_a.columns[3].width = Pt(50)

    # Set row height and font properties
    for row in new_table_a.rows:
        row.height = Pt(20)
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(11)

    for row_data in selected_questions_part_a:
        row_cells = new_table_a.add_row().cells
        for i, cell_value in enumerate(row_data):
            row_cells[i].text = cell_value

    # Add a paragraph break between Part-A and Part-B
    paragraph = new_doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run()
    run.add_break()

    # Add Part-B header
    part_b_heading = new_doc.add_paragraph("Part-B (2 x 15 = 30 Marks)", style='Heading1')
    part_b_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    part_b_heading.runs[0].bold = True
    part_b_heading.runs[0].font.size = Pt(14)

    # Add table for Part-B questions
    new_table_b = new_doc.add_table(rows=1, cols=4)
    new_table_b.style = 'Table Grid'  # Apply table style
    hdr_cells_b = new_table_b.rows[0].cells
    hdr_cells_b[0].text = "Q.No"  # Add Question Number header
    hdr_cells_b[1].text = "Questions"  # Add Questions header
    hdr_cells_b[2].text = "CO’s"  # Add CO's header
    hdr_cells_b[3].text = "Bloom’s Level"  # Add Bloom's Level header

    # Set column width
    new_table_b.columns[1].width = Pt(500)
    new_table_b.columns[0].width = Pt(50)
    new_table_b.columns[2].width = Pt(50)
    new_table_b.columns[3].width = Pt(50)

    # Set row height and font properties
    for row in new_table_b.rows:
        row.height = Pt(20)
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(11)

    for row_data in selected_questions_part_b:
        row_cells = new_table_b.add_row().cells
        for i, cell_value in enumerate(row_data):
            row_cells[i].text = cell_value

    # Save the new document
    new_doc.save(output_docx)
    print(f"Random questions saved to '{output_docx}'")
    fill_qno_ca1("output_questions.docx")


# ca2_paper("Format.docx", "output_questions.docx", "P123", "Physics", "PHY101", 3, "March 2024", "3rd")
# sem_paper("Format.docx", "output_questions.docx", "P123", "Physics", "PHY101","March 2024", "3rd")








